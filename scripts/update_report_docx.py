from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX_DIR = Path(r"D:\DATN\DOCX")
SOURCE = next(p for p in DOCX_DIR.glob("*.docx") if "THUY" in p.name.upper())
TARGET = DOCX_DIR / "Bao_cao_thuyet_minh_cap_nhat.docx"
SCREEN = ROOT / "assets" / "presentation" / "screens"
DOC_ASSETS = ROOT / "output" / "docx-assets"


def find_paragraph(doc, predicate):
    for para in doc.paragraphs:
        if predicate(para.text.strip()):
            return para
    raise ValueError("Paragraph not found")


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def replace_text(paragraph, text):
    paragraph.text = text
    return paragraph


def make_erd_image(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 34)
        box_font = ImageFont.truetype("arial.ttf", 24)
        small_font = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        title_font = ImageFont.load_default()
        box_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    draw.text((40, 30), "ERD Tong Quan - Junkio Expense Tracker", fill="black", font=title_font)

    boxes = {
        "users": (80, 190, 340, 290, (224, 242, 254)),
        "wallets": (80, 360, 340, 460, (204, 251, 241)),
        "categories": (400, 120, 700, 220, (254, 243, 199)),
        "transactions": (400, 300, 700, 420, (14, 165, 233)),
        "transaction_shares": (400, 520, 700, 640, (254, 226, 226)),
        "families": (850, 170, 1120, 270, (254, 243, 199)),
        "budgets": (850, 350, 1120, 450, (224, 242, 254)),
        "goals": (1200, 120, 1460, 220, (204, 251, 241)),
        "notifications": (1200, 300, 1460, 400, (224, 242, 254)),
        "audit_logs": (1200, 480, 1460, 580, (254, 226, 226)),
    }

    def box(name, fields):
        x1, y1, x2, y2, fill = boxes[name]
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline="black", width=2)
        draw.text((x1 + 16, y1 + 14), name, fill="black", font=box_font)
        draw.text((x1 + 16, y1 + 56), fields, fill="black", font=small_font)

    box("users", "id, email, role")
    box("wallets", "user_id, family_id, balance")
    box("categories", "user_id, type, name")
    box("transactions", "wallet_id, category_id, amount")
    box("transaction_shares", "transaction_id, user_id, status")
    box("families", "owner_id, name")
    box("budgets", "scope_type, scope_id, period")
    box("goals", "user_id, target_amount")
    box("notifications", "user_id, reference_id")
    box("audit_logs", "user_id, action, resource")

    def arrow(start, end):
        draw.line([start, end], fill="black", width=3)

    arrow((340, 240), (400, 350))
    arrow((340, 410), (400, 350))
    arrow((700, 350), (850, 220))
    arrow((700, 350), (850, 400))
    arrow((700, 350), (1200, 350))
    arrow((700, 580), (1200, 530))
    arrow((1120, 220), (1200, 170))
    arrow((1120, 220), (1200, 350))
    arrow((340, 240), (850, 220))
    arrow((340, 240), (1200, 530))

    draw.text((60, 770), "Users -> Wallets -> Transactions la truc du lieu chinh.", fill="black", font=small_font)
    draw.text((60, 805), "Families/transaction_shares mo rong nghiep vu chi tieu chung; budgets/goals/notifications/audit_logs ho tro dieu hanh va theo doi.", fill="black", font=small_font)
    img.save(path)


def add_picture_after(paragraph, picture_path, width_inches=5.8):
    pic_para = insert_paragraph_after(paragraph)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(picture_path), width=Inches(width_inches))
    return pic_para


def update_table_qa(doc):
    qa_table = next(t for t in doc.tables if t.cell(0, 0).text.strip() == "Hạng mục")
    rows = [
        ("Lint", "PASS", "ESLint pass; chỉ còn cảnh báo deopt trên font asset nhúng."),
        ("Unit / Integration test", "PASS", "Backend: 20 suite / 112 test. Frontend: 8 file test, 19 test pass, 2 skipped."),
        ("i18n", "PASS", "Locale parity và hardcoded text audit đạt yêu cầu."),
        ("Build", "PASS with warnings", "Build frontend thành công; còn cảnh báo chunk lớn ở xlsx, jspdf và font bundle."),
        ("Docker health", "PASS", "db, redis, api, web chạy ổn định; api, db, redis đều healthy."),
        ("Backend Newman", "PASS", "9 request, 17 assertion, 0 failure."),
        ("Admin API check", "PASS", "Lệnh check:admin trả về OK."),
        ("E2E desktop/mobile", "PASS", "Playwright smoke desktop và mobile đều pass; aggregate QA gate pass."),
    ]
    while len(qa_table.rows) - 1 < len(rows):
        qa_table.add_row()
    for idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            qa_table.cell(idx, col_idx).text = value


def main():
    DOC_ASSETS.mkdir(parents=True, exist_ok=True)
    erd_img = DOC_ASSETS / "erd_overview.png"
    make_erd_image(erd_img)

    doc = Document(str(SOURCE))

    flow_map = {
        "Auth": "Auth: Người dùng đăng ký hoặc đăng nhập để nhận access token JWT; refresh token được giữ qua cookie httpOnly nhằm tăng độ an toàn cho phiên làm việc.",
        "Transaction + Wallet": "Transaction + Wallet: Người dùng tạo ví, chọn danh mục và ghi nhận giao dịch; backend kiểm tra ownership rồi cập nhật số dư, thống kê và dữ liệu liên quan.",
        "Budget Alert": "Budget Alert: Khi mức chi đạt ngưỡng hoặc vượt hạn mức, hệ thống tạo notification, chống trùng theo ngày và có thể phát realtime đến giao diện.",
        "Recurring": "Recurring: Cron kiểm tra các recurring pattern đến hạn mỗi ngày, chỉ xử lý tối đa một kỳ overdue cho mỗi pattern để tránh sinh trùng hàng loạt.",
        "Family Shared Expense / Debt": "Family Shared Expense / Debt: Khoản chi chung được tách thành transaction_shares theo từng thành viên; từ đó hệ thống hỗ trợ approve/reject share, simplify debt và settle debt.",
        "Transfer": "Transfer: Một nghiệp vụ chuyển khoản tạo cặp bản ghi TRANSFER_OUT và TRANSFER_IN cùng transfer_group_id để bảo đảm truy vết và nhất quán dữ liệu.",
        "Export Report": "Export Report: Người dùng có thể xuất dữ liệu theo tiêu chí lọc ra CSV, Excel hoặc PDF để phục vụ thống kê, lưu trữ và in ấn.",
    }
    for old, new in flow_map.items():
        replace_text(find_paragraph(doc, lambda txt, old=old: txt == old), new)

    replace_text(find_paragraph(doc, lambda txt: txt == "[Chèn hình minh họa / sơ đồ tại đây]"), "Sơ đồ ERD dưới đây mô tả các thực thể lõi và mối liên hệ dữ liệu quan trọng nhất của hệ thống Junkio.")
    erd_anchor = find_paragraph(doc, lambda txt: txt.startswith("Sơ đồ ERD dưới đây mô tả"))
    add_picture_after(erd_anchor, erd_img, 6.6)
    replace_text(find_paragraph(doc, lambda txt: txt == "Hình x.x. 4.2. Sơ đồ ERD"), "Hình 4.1. Sơ đồ ERD của hệ thống")

    heading_51 = find_paragraph(doc, lambda txt: txt == "5.1. Sơ đồ luồng dữ liệu tổng quát")
    insert_paragraph_after(heading_51, "Luồng dữ liệu tổng quát của hệ thống được mô tả theo chu trình: người dùng thao tác trên giao diện web, frontend gửi request đến backend API, backend kiểm tra xác thực/quyền hạn rồi ghi nhận dữ liệu xuống PostgreSQL, sau đó phản hồi lại giao diện và phát thông báo realtime qua Socket.io khi cần.")

    transfer_req = find_paragraph(doc, lambda txt: txt == 'Request: {"fromWalletId": 1, "toWalletId": 2, "amount": 500000, "note": "Chuyển sang ví ngân hàng"}')
    insert_paragraph_after(transfer_req, 'Response: {"success": true, "data": {"transferGroupId": "grp_001", "outTransactionId": 201, "inTransactionId": 202, "amount": 500000}}')

    replace_text(find_paragraph(doc, lambda txt: txt.startswith("Phần backend được kiểm thử bằng Jest + Supertest")), "Phần backend được kiểm thử bằng Jest + Supertest cho các nghiệp vụ quan trọng như auth, wallets, transactions, budgets, families và admin flows. Ở lần kiểm tra gần nhất, backend smoke test đạt yêu cầu và tổng thể ghi nhận 20 suite / 112 test chạy thành công.")
    replace_text(find_paragraph(doc, lambda txt: txt.startswith("Frontend được kiểm thử bằng Vitest + Testing Library")), "Frontend được kiểm thử bằng Vitest + Testing Library ở các lớp hiển thị, biểu mẫu, điều hướng, phản hồi trạng thái tải dữ liệu, xử lý lỗi và trải nghiệm trên các thành phần dùng lại. Snapshot QA gần nhất ghi nhận 8 file test pass, 19 test pass và 2 test được skip có chủ đích.")
    replace_text(find_paragraph(doc, lambda txt: txt.startswith("API được kiểm thử bằng Postman/Newman")), "API được kiểm thử bằng Postman/Newman với các bộ test bao gồm trường hợp thành công, lỗi xác thực, dữ liệu không hợp lệ, dữ liệu thiếu trường bắt buộc và truy cập trái quyền. Đợt chạy gần nhất ghi nhận 9 request, 17 assertion và 0 failure.")

    update_table_qa(doc)

    demo_steps = [
        "Bước 1: Đăng nhập hệ thống bằng tài khoản demo hợp lệ để nhận phiên làm việc.",
        "Bước 2: Tạo ví mới hoặc mở ví có sẵn để chuẩn bị dữ liệu tài chính.",
        "Bước 3: Tạo giao dịch thu/chi và kiểm tra danh sách giao dịch cập nhật đúng.",
        "Bước 4: Tạo transfer giữa hai ví để chứng minh cặp TRANSFER_OUT / TRANSFER_IN.",
        "Bước 5: Tạo budget và quan sát cảnh báo khi chi tiêu chạm ngưỡng.",
        "Bước 6: Thiết lập recurring transaction và mô tả cơ chế cron xử lý định kỳ.",
        "Bước 7: Tạo family shared expense / debt để chứng minh nghiệp vụ chi tiêu chung.",
        "Bước 8: Xuất report CSV/Excel/PDF và mở Admin Dashboard để xem analytics.",
    ]
    for old, new in zip([
        "Đăng nhập hệ thống",
        "Tạo ví",
        "Tạo giao dịch",
        "Tạo transfer giữa hai ví",
        "Tạo budget và quan sát cảnh báo",
        "Thiết lập recurring transaction",
        "Tạo family shared expense / debt",
        "Xuất report CSV/Excel/PDF",
    ], demo_steps):
        replace_text(find_paragraph(doc, lambda txt, old=old: txt == old), new)
    replace_text(find_paragraph(doc, lambda txt: txt == "Mở Admin Dashboard để theo dõi analytics và người dùng"), "Bước 9: Mở Admin Dashboard để theo dõi analytics, người dùng và log quản trị.")

    replace_text(find_paragraph(doc, lambda txt: txt.startswith("Phần demo giao diện nên trình bày theo luồng thao tác thực tế")), "Phần demo giao diện được minh họa bằng ảnh chụp thật từ hệ thống đang chạy. Các màn hình Dashboard, Transactions, Family, Goals và Admin Dashboard phản ánh trực tiếp trạng thái hiện tại của sản phẩm, giúp phần báo cáo không còn phụ thuộc vào placeholder mô tả.")

    for old, new in {
        "• Giao diện đăng nhập": "• Giao diện đăng nhập: minh chứng cho luồng xác thực người dùng và điều hướng đầu vào.",
        "• Giao diện dashboard tổng quan": "• Giao diện dashboard tổng quan: hiển thị số liệu tài chính, biểu đồ và các chỉ số chính.",
        "• Giao diện quản lý giao dịch": "• Giao diện quản lý giao dịch: hỗ trợ tìm kiếm, lọc, phân trang và xem chi tiết.",
        "• Giao diện quản lý ngân sách": "• Giao diện quản lý ngân sách: thể hiện hạn mức, tiến độ sử dụng và cảnh báo.",
        "• Giao diện gia đình và công nợ": "• Giao diện gia đình và công nợ: minh họa chi tiêu chung, chia sẻ khoản chi và nghĩa vụ thanh toán.",
    }.items():
        replace_text(find_paragraph(doc, lambda txt, old=old: txt == old), new)

    append_b = find_paragraph(doc, lambda txt: txt.startswith("• Giao diện gia đình và công nợ"))
    for picture in [SCREEN / "dashboard.png", SCREEN / "transactions.png", SCREEN / "admin.png"]:
        append_b = add_picture_after(append_b, picture, 5.9)

    appendix_d = find_paragraph(doc, lambda txt: txt == "D. Kịch bản kiểm thử")
    insert_paragraph_after(appendix_d, "Các kịch bản trong bảng D.1 được dùng để đối chiếu giữa kiểm thử tự động và demo thủ công, bao phủ các nhóm auth, transaction, recurring, budget alert, debt share, export và admin.")

    replace_text(find_paragraph(doc, lambda txt: txt == "Phát triển mobile app riêng"), "Phát triển mobile app hoặc PWA riêng để tăng tính cơ động và tối ưu thao tác trên thiết bị di động.")
    replace_text(find_paragraph(doc, lambda txt: txt == "Tích hợp ngân hàng / ví điện tử"), "Tích hợp ngân hàng hoặc ví điện tử nhằm giảm thao tác nhập liệu thủ công và tăng tính thực tiễn của sản phẩm.")
    replace_text(find_paragraph(doc, lambda txt: txt == "Nâng cấp forecast bằng AI/ML"), "Nâng cấp forecast từ mức thống kê cơ bản lên mô hình AI/ML để gợi ý ngân sách và dự đoán chi tiêu thông minh hơn.")
    replace_text(find_paragraph(doc, lambda txt: txt == "Bổ sung alert/rule nâng cao"), "Bổ sung alert/rule nâng cao để cảnh báo bất thường theo danh mục, khoảng thời gian hoặc hành vi chi tiêu.")
    replace_text(find_paragraph(doc, lambda txt: txt == "Hoàn thiện CI/CD cho production"), "Hoàn thiện CI/CD cho môi trường production để tăng mức độ tự động hóa triển khai và kiểm soát chất lượng.")
    replace_text(find_paragraph(doc, lambda txt: txt == "Bổ sung monitoring và backup dữ liệu tự động"), "Bổ sung monitoring, logging tập trung và backup dữ liệu tự động để tăng độ sẵn sàng khi vận hành thực tế.")

    e_anchor = find_paragraph(doc, lambda txt: txt == "Các bước thực hiện:")
    insert_paragraph_after(e_anchor, "Lưu ý: nếu cần dựng lại môi trường sạch hoàn toàn, có thể chạy docker compose down -v --remove-orphans trước khi docker compose up --build -d.")

    doc.save(str(TARGET))
    print(TARGET)


if __name__ == "__main__":
    main()
